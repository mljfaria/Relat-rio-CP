from flask import Flask, render_template, request, send_file
from docx import Document
import os

app = Flask(__name__)

# Criar a pasta static/ se não existir
if not os.path.exists("static"):
    os.makedirs("static")

# Verificar se o modelo de relatório existe
modelo_path = "static/modelo_relatorio.docx"
if not os.path.exists(modelo_path):
    with open(modelo_path, 'w') as f:
        f.write("Este é um modelo temporário. Substitua-o pelo modelo real.")

# Página principal com formulário
@app.route('/')
def index():
    return '''
    <html>
        <body>
            <h2>Preencher Relatório</h2>
            <form action="/gerar_word" method="post">
                Data: <input type="text" name="data"><br>
                Indicativo: <input type="text" name="indicativo"><br>
                Turno: <input type="text" name="turno"><br>
                Chefe de Viatura Nome: <input type="text" name="chefe_nome"><br>
                Chefe de Viatura Posto: <input type="text" name="chefe_posto"><br>
                Condutor Nome: <input type="text" name="condutor_nome"><br>
                Condutor Posto: <input type="text" name="condutor_posto"><br>
                Viatura Matrícula: <input type="text" name="viatura_matricula"><br>
                Viatura Modelo: <input type="text" name="viatura_modelo"><br>
                Ocorrências: <textarea name="ocorrencias"></textarea><br>
                <input type="submit" value="Gerar Relatório">
            </form>
        </body>
    </html>
    '''

# Geração do relatório em Word
@app.route('/gerar_word', methods=['POST'])
def gerar_word():
    # Capturar dados do formulário
    novo_doc_path = "static/Relatorio_Preenchido.docx"
    
    if not os.path.exists(modelo_path):
        return "Erro: O modelo de relatório Word não foi encontrado. Certifique-se de que o ficheiro modelo_relatorio.docx está na pasta static."
    
    doc = Document(modelo_path)
    
    dados_relatorio = {
        "Data": request.form['data'],
        "Indicativo": request.form['indicativo'],
        "Turno": request.form['turno'],
        "Chefe de Viatura": request.form['chefe_nome'],
        "Posto Chefe": request.form['chefe_posto'],
        "Condutor": request.form['condutor_nome'],
        "Posto Condutor": request.form['condutor_posto'],
        "Viatura Matrícula": request.form['viatura_matricula'],
        "Viatura Modelo": request.form['viatura_modelo'],
        "Ocorrências": request.form['ocorrencias']
    }
    
    for paragrafo in doc.paragraphs:
        for chave, valor in dados_relatorio.items():
            if chave in paragrafo.text:
                paragrafo.text = paragrafo.text.replace(chave, valor)
    
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for chave, valor in dados_relatorio.items():
                    if chave in celula.text:
                        celula.text = celula.text.replace(chave, valor)
    
    doc.save(novo_doc_path)
    
    return send_file(novo_doc_path, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
